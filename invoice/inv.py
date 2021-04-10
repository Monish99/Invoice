from docx import Document
import mysql
from mysql.connector import Error
from docx2pdf import convert
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_connection(hn, un, pw, dbn):
    connection = None
    try:
        connection = mysql.connector.connect(
            host=hn,
            user=un,
            passwd=pw,
            database=dbn
        )
        print("MySQL Database connection successful")
    except Error as err:
        print(f"Error: '{err}'")
    return connection


def read_query(con, que):
    cursor = con.cursor()
    result = None
    try:
        cursor.execute(que)
        result = cursor.fetchall()
        return result
    except Error as err:
        print(f"Error:", err)


def someRandom(tr, col, mode, noDays, sd):
    print("type of triner ", type(tr))
    conn = create_connection("localhost", "root", "", "genesis")
    # Fetching all details of trainer
    que = "SELECT * FROM trainer where Name='" + tr + "'"
    que1 = "SELECT BName FROM trainer where Name='" + tr + "'"
    TBD = str(read_query(conn, que))
    # bn = str(read_query(conn, que1))
    print(TBD)
    l = TBD.split("'")
    for i in range(len(l)):
        print(i, "vlaue ; ", l[i])
    # Fetching trainer and college lacation for cal travel allowance
    que = "SELECT Location FROM collage where Name='" + col + "'"
    que1 = "SELECT Location FROM trainer where Name='" + tr + "'"
    colLoc = str(read_query(conn, que))
    trLoc = str(read_query(conn, que1))
    colLoc1 = colLoc.split("'")
    trLoc1 = trLoc.split("'")
    colLoc = colLoc1[1]
    trLoc = trLoc1[1]
    print("col ; ", colLoc, "trloc ; ", trLoc)
    # Fetching food details from DB
    que1 = "SELECT Food FROM collage where Name='" + col + "'"
    fo = str(read_query(conn, que1))
    fo1 = fo.split("'")
    fo = fo1[1]
    print("foo : ", fo, "type", type(fo))
    # Selecting Pay of trainer
    que1 = "SELECT Payment FROM trainer where Name='" + tr + "'"
    pa = str(read_query(conn, que1))
    print("PA : ", pa)
    pa1 = pa.split("'")
    pa2 = pa1[0].split("(")
    print("pa2  ", pa2)
    pa1 = pa2[1].split(",")
    print("pa1 ; ", pa1)
    pa = pa1[0]
    print("Start date", sd)

    invGen(l, trLoc, colLoc, mode, fo, noDays, pa, sd, col)


def invGen(lst1, TLoc, CLoc, mode, foo, NoD, pay, stDate, col):
    doc = Document()
    my_image = doc.add_picture('invoice/some.png')
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    Phone = lst1[2].split(',')
    print("aa : ", Phone)
    # Adding Heading to docx
    doc.add_heading('Genplus Training and Consulting', 0)
    # Defining the dynamic properties that have to be printed in the docx like trainer name etc..
    name = "Name(As given in Bank) : " + lst1[1]
    bacc = "Bank Account Number : " + lst1[9]
    ifsc = "IFSC : " + lst1[11]
    pan = "PAN Number : " + lst1[13]
    bn = "Bank Name : " + lst1[15]
    ph = "Phone Number : " + Phone[1]
    em = "Email : " + lst1[3]
    loc = "Based Location : " + lst1[5]
    # Adding each of the dynamic functionality like name
    para = doc.add_paragraph()
    run2 = para.add_run()
    run2.add_text(name)
    # Adding Acc no
    para = doc.add_paragraph()
    run2 = para.add_run()
    run2.add_text(bacc)
    # Adding IFSC code
    para = doc.add_paragraph()
    run2 = para.add_run()
    run2.add_text(ifsc)
    # Adding PAN
    para = doc.add_paragraph()
    run2 = para.add_run()
    run2.add_text(pan)
    # Adding Bank Name
    para = doc.add_paragraph()
    run2 = para.add_run()
    run2.add_text(bn)

    para = doc.add_paragraph()
    run2 = para.add_run()
    run2.add_text(ph)

    para = doc.add_paragraph()
    run2 = para.add_run()
    run2.add_text(em)

    para = doc.add_paragraph()
    run2 = para.add_run()
    run2.add_text(loc)


    # Calculating Food and travel Allowance
    food = "No"
    travel = "No"
    if mode == "offline":
        if TLoc != CLoc:
            travel = "1000"
        else:
            travel = "No"
        if foo == "No":
            food = '200'
        else:
            food = 'No'
    # Formatting start_date as a str
    stDate1 = stDate.split("-")  # strdate[y, mon, day]
    day = int(stDate1[2])
    mon = stDate1[1]
    yr = stDate1[0]
    nod = int(NoD)
    # Assigning all the data to be put into the table to a Tuple
    data = []
    for i in range(nod + 1):
        date = str(day) + "/" + mon + "/" + yr
        data1 = (date, col, pay, travel, food)
        data.append(data1)
        day = day + 1
    fday = day
    print()
    da = tuple(data)
    table = doc.add_table(rows=1, cols=5)
    # Adding heading in the 1st row of the table
    row = table.rows[0].cells
    row[0].text = 'Date'
    row[1].text = 'College'
    row[2].text = 'Fees/day'
    row[3].text = 'Travel Allowance'
    row[4].text = 'Food Allowance'
    # Adding data from the list to the table
    day = int(stDate1[2])
    print('day : ', day, "stdate[2] ; ", stDate1[2], "Fday ; ", fday)
    for d, c, f, t, fo in data:
        if t == "1000":
            if day == int(stDate1[2]) or day == fday - 1:
                t = "1000"
            else:
                t = 'No'
        row = table.add_row().cells
        row[0].text = d
        row[1].text = c
        row[2].text = f
        row[3].text = t
        row[4].text = fo
        day = day + 1
    # Empty row
    row = table.add_row().cells
    row = table.add_row().cells
    # Calculating Total amount of each col
    tot_fee = int(pay) * (int(NoD) + 1)
    row[1].text = "Total"
    row[2].text = str(tot_fee)
    tot_travel = 0
    if travel == '1000':
        row[3].text = '2000'
        tot_travel = 2000
    tot_food = 0
    if food == '200':
        tot_food = 200 * (int(NoD) + 1)
        row[4].text = str(tot_food)
    row = table.add_row().cells
    row = table.add_row().cells
    # Calculating the Grand_total
    G_tot = tot_fee + tot_food + tot_travel
    row[3].text = "Grand Total"
    row[4].text = str(G_tot)
    # Styling the table
    table.style = 'Light Grid Accent 1'
    filePath = 'C:\\Users\\Monish\\invoiceGen\\Team_Void\\invoice.docx'
    doc.save(filePath)
    #print("This is random : ")
    convert(filePath)

# someRandom('Monish', 'SJBIT', "offline", '6', "2021-04-05")
